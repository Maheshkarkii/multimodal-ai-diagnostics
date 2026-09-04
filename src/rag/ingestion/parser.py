"""
Document Ingestion and Parsing Pipeline.
Supports PDF, TXT, Markdown, and DOCX files.
Preserves page numbers, computes cryptographic content hashes, and flags scanned/empty pages.
"""

import hashlib
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

from src.rag.config import DocumentIngestionConfig
from src.rag.schema import DocumentMetadata, RawDocumentPage

logger = logging.getLogger(__name__)


def compute_file_hash(file_path: Union[str, Path]) -> str:
    """Compute SHA-256 hash of a file for incremental indexing detection."""
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        while chunk := f.read(65536):
            sha256.update(chunk)
    return sha256.hexdigest()


class DocumentParser:
    """Base class and unified dispatcher for extracting structured text from technical documents."""

    def __init__(self, config: Optional[DocumentIngestionConfig] = None):
        self.config = config or DocumentIngestionConfig()

    def parse_document(
        self,
        file_path: Union[str, Path],
        extra_metadata: Optional[Dict] = None
    ) -> Tuple[DocumentMetadata, List[RawDocumentPage]]:
        """Parse any supported technical document into DocumentMetadata and a list of RawDocumentPage objects."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {path}")

        file_ext = path.suffix.lower()
        if file_ext not in self.config.supported_extensions:
            raise ValueError(
                f"Unsupported document format: '{file_ext}'. Supported: {self.config.supported_extensions}"
            )

        file_hash = compute_file_hash(path)
        file_size = path.stat().st_size
        doc_id = f"{path.stem}_{file_hash[:10]}"
        doc_name = path.name

        extra_meta = extra_metadata or {}
        # Parse based on extension
        if file_ext == ".pdf":
            pages = self._parse_pdf(path)
        elif file_ext in [".txt", ".md"]:
            pages = self._parse_text_or_markdown(path)
        elif file_ext == ".docx":
            pages = self._parse_docx(path)
        else:
            raise ValueError(f"No parser handler registered for: {file_ext}")

        # Construct metadata
        metadata = DocumentMetadata(
            document_id=doc_id,
            document_name=doc_name,
            source_path=str(path.resolve()),
            file_type=file_ext.lstrip("."),
            file_hash=file_hash,
            file_size_bytes=file_size,
            num_pages=len(pages),
            equipment_type=extra_meta.get("equipment_type"),
            manufacturer=extra_meta.get("manufacturer"),
            model=extra_meta.get("model"),
            manual_version=extra_meta.get("manual_version"),
            publication_date=extra_meta.get("publication_date"),
            custom_metadata=extra_meta.get("custom_metadata", {}),
        )

        return metadata, pages

    def _parse_pdf(self, path: Path) -> List[RawDocumentPage]:
        """Parse PDF document with page boundary preservation and scanned-text quality detection."""
        pages: List[RawDocumentPage] = []

        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))

            for page_idx, page in enumerate(reader.pages, start=1):
                raw_text = page.extract_text() or ""
                cleaned_text = self._clean_text(raw_text)

                # Detect poor extraction or scanned pages
                is_scanned = False
                if self.config.detect_scanned_pdf:
                    if len(cleaned_text.strip()) < self.config.min_page_chars_threshold:
                        is_scanned = True
                        logger.warning(
                            f"Page {page_idx} of '{path.name}' contains fewer than {self.config.min_page_chars_threshold} "
                            f"characters ({len(cleaned_text.strip())} chars). Marked as potentially scanned/sparse."
                        )

                # Detect header/section heuristic
                section_title = self._detect_section_title(cleaned_text)
                has_tables = self._detect_table_structure(cleaned_text)

                pages.append(
                    RawDocumentPage(
                        page_number=page_idx,
                        text=cleaned_text,
                        section_title=section_title,
                        has_tables=has_tables,
                        is_scanned_suspicious=is_scanned,
                    )
                )

        except Exception as e:
            logger.error(f"Error parsing PDF '{path}': {e}")
            raise RuntimeError(f"Failed to extract PDF content from {path}: {e}") from e

        return pages

    def _parse_text_or_markdown(self, path: Path) -> List[RawDocumentPage]:
        """Parse TXT or Markdown files. Simulates pages using page markers or logical sections."""
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        cleaned_content = self._clean_text(content)

        # Look for explicit page breaks like "--- PAGE 2 ---" or Markdown horizontal rules with page info
        raw_sections = cleaned_content.split("\n--- PAGE ")
        pages: List[RawDocumentPage] = []

        if len(raw_sections) > 1:
            # First segment before any explicit page marker
            p1_text = raw_sections[0].strip()
            pages.append(
                RawDocumentPage(
                    page_number=1,
                    text=p1_text,
                    section_title=self._detect_section_title(p1_text),
                    has_tables=self._detect_table_structure(p1_text),
                )
            )
            for idx, sec in enumerate(raw_sections[1:], start=2):
                # May start with "2 ---\ntext..."
                lines = sec.split("\n", 1)
                text = lines[1].strip() if len(lines) > 1 else sec.strip()
                pages.append(
                    RawDocumentPage(
                        page_number=idx,
                        text=text,
                        section_title=self._detect_section_title(text),
                        has_tables=self._detect_table_structure(text),
                    )
                )
        else:
            # Single page document or large markdown file split logically
            pages.append(
                RawDocumentPage(
                    page_number=1,
                    text=cleaned_content,
                    section_title=self._detect_section_title(cleaned_content),
                    has_tables=self._detect_table_structure(cleaned_content),
                )
            )

        return pages

    def _parse_docx(self, path: Path) -> List[RawDocumentPage]:
        """Parse Microsoft Word DOCX files if python-docx is present, otherwise fallback gracefully."""
        try:
            import docx
            doc = docx.Document(str(path))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            combined = "\n\n".join(full_text)
            return [
                RawDocumentPage(
                    page_number=1,
                    text=self._clean_text(combined),
                    section_title=self._detect_section_title(combined),
                    has_tables=len(doc.tables) > 0,
                )
            ]
        except ImportError:
            logger.warning("python-docx is not installed. Skipping DOCX parsing or treat as unsupported.")
            raise NotImplementedError("python-docx package required for parsing .docx files.")

    def _clean_text(self, text: str) -> str:
        """Clean extracted document text, normalize whitespace, and remove null bytes."""
        if not self.config.clean_whitespace:
            return text

        # Strip null bytes & control chars
        text = text.replace("\x00", " ")
        # Normalize carriage returns
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Collapse excessive blank lines
        lines = [line.strip() for line in text.split("\n")]
        cleaned_lines = []
        consecutive_blank = 0
        for line in lines:
            if not line:
                consecutive_blank += 1
                if consecutive_blank <= 1:
                    cleaned_lines.append("")
            else:
                consecutive_blank = 0
                cleaned_lines.append(line)

        return "\n".join(cleaned_lines).strip()

    def _detect_section_title(self, text: str) -> Optional[str]:
        """Heuristic to detect section headings (e.g., '1.0 Introduction', 'SECTION 3 - BEARING INSPECTION', '## Vibration Limits')."""
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        for line in lines[:5]:  # Look in first few lines
            if line.startswith("#"):
                return line.lstrip("#").strip()
            if any(line.upper().startswith(prefix) for prefix in ["SECTION", "CHAPTER", "PART", "PROCEDURE", "MODULE"]):
                return line
            # Numbered headings like '1.2 Vibration Thresholds'
            if len(line) > 3 and line[0].isdigit() and ("." in line[:4] or " " in line[:4]):
                return line
        return None

    def _detect_table_structure(self, text: str) -> bool:
        """Check if page contains table formatting indicators such as pipe '|' columns or tab separations."""
        lines = text.split("\n")
        pipe_lines = sum(1 for line in lines if line.count("|") >= 2)
        tab_lines = sum(1 for line in lines if line.count("\t") >= 2)
        return pipe_lines >= 2 or tab_lines >= 2
